from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
import uuid
import uvicorn

app = FastAPI(title="Smart Exam API Only")

# BẮT BUỘC PHẢI CÓ: Cho phép file HTML từ máy bạn kết nối tới Server Online
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = "/tmp" if os.path.exists("/tmp") else "."

# --- CÁC HÀM XỬ LÝ LOGIC (GIỮ NGUYÊN) ---
def extract_options_smart(text):
    """Tách các lựa chọn A, B, C, D từ dòng văn bản"""
    pattern = r"([A-D]|[a-d])[\.\)\:]\s+"
    parts = re.split(pattern, text)
    options = []
    if len(parts) > 2:
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                label = parts[i].upper()
                content = parts[i+1].strip()
                options.append({"label": label, "content": content})
    return options
def parse_answer_zone(paragraphs):
    """
    Phân tích vùng đáp án để lấy:
    1. Đáp án đúng (A, B, C, D hoặc trắc nghiệm đúng sai)
    2. Lời giải thích (nếu có)
    """
    answer_map = {}
    explanation_map = {} # Lưu lời giải
    
    current_id = None
    p_q_id = re.compile(r"^(Câu\s+(\d+))[:\.]?", re.IGNORECASE)
    
    # Regex bắt dòng đáp án: "Đáp án: A" hoặc "Đáp án: a) S, b) Đ"
    p_ans_line = re.compile(r"[:\.\·\-\s]*Đáp án\s*[:\.]?\s*(.*)", re.IGNORECASE)
    
    # Regex bắt đầu lời giải
    p_explain_start = re.compile(r"^(Giải thích|Hướng dẫn|Lời giải)[:\.]?", re.IGNORECASE)
    
    current_explanation = []
    is_collecting_explanation = False

    for para in paragraphs:
        text = para.text.strip()
        if not text: continue

        # 1. Phát hiện Câu mới trong vùng đáp án
        match_id = p_q_id.match(text)
        if match_id:
            # Lưu lời giải cũ nếu đang thu thập
            if current_id and current_explanation:
                explanation_map[current_id] = "\n".join(current_explanation).strip()
                current_explanation = []
            
            current_id = match_id.group(2)
            is_collecting_explanation = False # Reset trạng thái
            
            # Nếu dòng này chứa luôn "Đáp án:", xử lý ngay
            if "Đáp án" in text: 
                pass # Để logic phía dưới xử lý tiếp phần text
            else:
                continue

        if current_id:
            # 2. Xử lý dòng Đáp án
            match_ans = p_ans_line.search(text)
            if match_ans:
                ans_content = match_ans.group(1).strip() # Giữ nguyên case để check a) b)
                ans_upper = ans_content.upper()
                
                # Checkbox (Đúng/Sai): a) S, b) Đ...
                # Regex tìm a) Đ hoặc A. ĐÚNG
                checkbox_matches = re.findall(r"([A-Da-d])[\)\.\:]\s*(?:Đ|TRUE|ĐÚNG|S|FALSE|SAI)", ans_content, re.IGNORECASE)
                
                if checkbox_matches:
                    indices = []
                    mapping = {'A': 1, 'B': 2, 'C': 3, 'D': 4, 'a': 1, 'b': 2, 'c': 3, 'd': 4}
                    
                    # Logic phức tạp hơn: Phải check xem nó là Đ hay S. 
                    # Nhưng Quizizz thường chỉ cần biết câu nào là Đáp án đúng.
                    # Với dạng Đúng/Sai của Quizizz, ta thường không import được trực tiếp logic từng ý.
                    # Ở đây ta tạm lấy mapping index để phục vụ tô màu Word.
                    for char in checkbox_matches:
                         if char.upper() in mapping: indices.append(mapping[char.upper()])
                    
                    # Cải tiến: Nếu là dạng a) S, b) Đ -> Ta cần biết ý nào là ĐÚNG (Đ).
                    # Regex tìm cụ thể những ý ĐÚNG
                    true_matches = re.findall(r"([A-Da-d])[\)\.\:]\s*(?:Đ|TRUE|ĐÚNG)", ans_content, re.IGNORECASE)
                    true_indices = []
                    for char in true_matches:
                        if char.upper() in mapping: true_indices.append(mapping[char.upper()])
                    
                    answer_map[current_id] = sorted(list(set(true_indices))) if true_indices else []
                    
                else:
                    # Trắc nghiệm thường: Đáp án: A
                    mc_match = re.search(r"\b([A-D])\b", ans_upper)
                    if mc_match:
                        char = mc_match.group(1)
                        mapping = {'A': 1, 'B': 2, 'C': 3, 'D': 4}
                        answer_map[current_id] = [mapping.get(char, 1)]
            
            # 3. Xử lý Lời giải
            if p_explain_start.match(text):
                is_collecting_explanation = True
                # Loại bỏ từ khóa "Giải thích:" khỏi nội dung
                content = re.sub(r"^(Giải thích|Hướng dẫn|Lời giải)[:\.]?\s*", "", text, flags=re.IGNORECASE)
                if content: current_explanation.append(content)
            elif is_collecting_explanation:
                # Nếu đang trong chế độ gom lời giải, và dòng này không phải Câu mới hay Đáp án
                if not match_ans:
                    current_explanation.append(text)
    
    # Lưu lời giải câu cuối cùng
    if current_id and current_explanation:
        explanation_map[current_id] = "\n".join(current_explanation).strip()

    return answer_map, explanation_map
def parse_docx_split_mode(doc):
    all_paras = doc.paragraphs
    split_index = -1
    for i, para in enumerate(all_paras):
        if any(kw in para.text.strip().upper() for kw in ["ĐÁP ÁN VÀ GIẢI THÍCH", "HƯỚNG DẪN CHẤM", "PHẦN ĐÁP ÁN"]):
            split_index = i; break
    question_paras, answer_paras = (all_paras[:split_index], all_paras[split_index:]) if split_index != -1 else (all_paras, [])
    key_map, explain_map = parse_answer_zone(answer_paras)
    questions, current_q, current_part = [], {}, 1
    p_header, p_q_start = re.compile(r"^PHẦN\s+(\d+)", re.IGNORECASE), re.compile(r"^(Câu\s+(\d+))[:\.]?\s*(.*)", re.IGNORECASE)
    def save_q():
        nonlocal current_q
        if current_q:
            q_id = current_q["id"]
            current_q["correct_indices"] = key_map.get(q_id, [])
            current_q["explanation"] = explain_map.get(q_id, "")
            questions.append(current_q)
        current_q = {}
    for para in question_paras:
        text = para.text.strip()
        if not text: continue
        m_header = p_header.match(text)
        if m_header: save_q(); current_part = int(m_header.group(1)); continue
        m_q = p_q_start.match(text)
        if m_q:
            save_q()
            current_q = {
                "id": m_q.group(2), "text": m_q.group(3).strip(), 
                "type": "Checkbox" if current_part == 2 else ("Open-Ended" if current_part == 3 else "Multiple Choice"),
                "time": 60 if current_part == 2 else (300 if current_part == 3 else 30),
                "options": [], "correct_indices": [], "explanation": ""
            }
            continue
        if not current_q: continue
        if current_part != 3:
            opts_inline = extract_options_smart(text)
            if opts_inline: current_q["options"].extend(opts_inline)
            else:
                m_opt = re.match(r"^([A-D]|[a-d])[\.\)\:]\s*(.*)", text)
                if m_opt: current_q["options"].append({"label": m_opt.group(1).upper(), "content": m_opt.group(2).strip()})
                else: 
                    if len(current_q["options"]) == 0: current_q["text"] += "\n" + text
        else: current_q["text"] += "\n" + text
    save_q()
    return questions

def export_excel(questions, file_path):
    rows = []
    columns = ["Question Text", "Question Type", "Option 1", "Option 2", "Option 3", "Option 4", "Option 5", "Correct Answer", "Time in seconds", "Explanation"]
    for q in questions:
        row = {col: "" for col in columns}
        row["Question Text"], row["Question Type"], row["Time in seconds"], row["Explanation"] = q["text"], q["type"], q["time"], q["explanation"]
        for i, opt in enumerate(q["options"][:5]): row[f"Option {i+1}"] = opt["content"]
        if q["type"] != "Open-Ended": row["Correct Answer"] = ",".join(map(str, q["correct_indices"])) if q["correct_indices"] else "1"
        rows.append(row)
    pd.DataFrame(rows, columns=columns).to_excel(file_path, index=False)

def export_word_bold(questions, file_path):
    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Times New Roman', Pt(12)
    doc.add_heading('ĐỀ THI (ĐÁP ÁN IN ĐẬM ĐỎ)', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for q in questions:
        p = doc.add_paragraph()
        p.add_run(f"Câu {q['id']}: ").bold = True
        p.add_run(q['text'])
        if q['options']:
            for i, opt in enumerate(q['options']):
                is_correct = (i + 1) in q['correct_indices']
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Pt(18)
                run_l = p_opt.add_run(f"{chr(97+i)+')' if q['type']=='Checkbox' else opt['label']+'.'} ")
                run_c = p_opt.add_run(f"{opt['content']}")
                if is_correct:
                    run_l.bold = run_c.bold = True
                    run_l.font.color.rgb = run_c.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(file_path)

def export_word_standard(questions):
    """
    Xuất file Word chuẩn 3 phần:
    1. Đề thi (chỉ có câu hỏi, options xếp dọc đẹp)
    2. Bảng đáp án (Grid)
    3. Lời giải chi tiết
    """
    doc = Document()
    # Cấu hình font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # --- PHẦN 1: ĐỀ THI ---
    doc.add_heading('ĐỀ KIỂM TRA', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Môn: Tin học | Thời gian: 45 phút").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("_" * 30).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for q in questions:
        # In câu hỏi
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_idx = p.add_run(f"Câu {q['id']}: ")
        run_idx.bold = True
        p.add_run(q['text'])
        
        # In options (Luôn xuống dòng cho đẹp - Standard Format)
        if q['options']:
            is_checkbox = (q['type'] == "Checkbox")
            for i, opt in enumerate(q['options']):
                label_char = opt['label']
                if is_checkbox: label_char = chr(97 + i) + ")" 
                else: label_char = label_char + "."
                
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Pt(24) # Thụt lề sâu hơn tí
                p_opt.paragraph_format.space_after = Pt(0) # Sát nhau
                p_opt.add_run(f"{label_char} {opt['content']}")
        
        doc.add_paragraph() # Dòng trống giữa các câu

    # Ngắt trang
    doc.add_page_break()

    # --- PHẦN 2: BẢNG ĐÁP ÁN ---
    doc.add_heading('BẢNG ĐÁP ÁN', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph() # Dòng trống
    
    # === SỬA LỖI TẠI ĐÂY ===
    # Thay vì tạo bảng rồi xóa, ta tính toán số dòng cần thiết TRƯỚC
    num_q = len(questions)
    cols = 5 # 5 câu trên 1 dòng
    
    if num_q > 0:
        # Tính số hàng cần thiết (làm tròn lên)
        rows_needed = (num_q // cols) + (1 if num_q % cols > 0 else 0)
        
        # Tạo bảng 1 lần duy nhất với kích thước chuẩn
        # rows_needed * 2 vì mỗi câu cần 2 dòng: 1 dòng ID câu hỏi, 1 dòng Đáp án
        table = doc.add_table(rows=rows_needed * 2, cols=cols)
        table.style = 'Table Grid'
        
        for idx, q in enumerate(questions):
            r = (idx // cols) * 2
            c = idx % cols
            
            # Ô ID câu hỏi
            cell_id = table.cell(r, c)
            cell_id.text = str(q['id'])
            cell_id.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Ô Đáp án
            cell_ans = table.cell(r + 1, c)
            
            ans_text = ""
            if q['correct_indices']:
                labels = ['A', 'B', 'C', 'D']
                if q['type'] == "Checkbox":
                    # Với Checkbox, hiển thị các ý ĐÚNG (ví dụ: a,c)
                    ans_text = ",".join([chr(96 + i) for i in q['correct_indices']])
                else:
                    # Trắc nghiệm thường
                    # Kiểm tra index có hợp lệ không để tránh lỗi list index out of range
                    if len(q['correct_indices']) > 0:
                        idx_ans = q['correct_indices'][0] - 1
                        if 0 <= idx_ans < 4:
                            ans_text = labels[idx_ans]
            
            cell_ans.text = ans_text
            cell_ans.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Chỉ format nếu có đáp án
            if ans_text:
                run = cell_ans.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        doc.add_paragraph("Không tìm thấy câu hỏi trắc nghiệm nào.")

    doc.add_paragraph()
    
    # --- PHẦN 3: LỜI GIẢI CHI TIẾT ---
    doc.add_heading('LỜI GIẢI CHI TIẾT', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    has_explanation = False
    for q in questions:
        if q.get('explanation') and str(q['explanation']).strip():
            has_explanation = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            
            run_h = p.add_run(f"Câu {q['id']}: ")
            run_h.bold = True
            run_h.italic = True
            run_h.font.color.rgb = RGBColor(0, 0, 255) # Màu xanh cho dễ nhìn
            
            p.add_run(q['explanation'])
            
    if not has_explanation:
        doc.add_paragraph("(Không có lời giải chi tiết được trích xuất)")
    doc.save(file_path)

# --- ENDPOINTS API ---
@app.post("/api/process")
async def process_exam_text(request: Request):
    try:
        data = await request.json()
        raw_text = data.get("text", "")
        if not raw_text: return {"error": "Nội dung trống!"}
        
        doc = Document()
        for line in raw_text.split('\n'):
            if line.strip(): doc.add_paragraph(line.strip())
            
        questions = parse_docx_split_mode(doc)
        uid = uuid.uuid4().hex[:8]
        
        excel_name = f"Quizizz_{uid}.xlsx"
        word_bold_name = f"WordBold_{uid}.docx"
        word_std_name = f"WordStd_{uid}.docx"
        
        export_excel(questions, os.path.join(BASE_DIR, excel_name))
        export_word_bold(questions, os.path.join(BASE_DIR, word_bold_name))
        export_word_standard(questions, os.path.join(BASE_DIR, word_std_name))
        
        return {
            "excel_url": f"/api/download/{excel_name}",
            "word_bold_url": f"/api/download/{word_bold_name}",
            "word_std_url": f"/api/download/{word_std_name}"
        }
    except Exception as e:
        return {"error": str(e)}

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(BASE_DIR, filename)
    if os.path.exists(file_path):
        friendly_name = "Quizizz_Import_File.xlsx"
        if "WordBold" in filename: friendly_name = "De_Thi_Bold_Dap_An.docx"
        elif "WordStd" in filename: friendly_name = "De_Thi_Chuan_Format_Full.docx"
        return FileResponse(file_path, media_type="application/octet-stream", filename=friendly_name)
    return {"error": "File không tồn tại"}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
